from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re

def create_word_document(input_txt, output_docx):
    doc = Document()

    # Configure document for RTL
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)
    # The python-docx RTL support is somewhat limited natively without modifying XML directly,
    # but we can set paragraph alignment to right.

    heading = doc.add_heading('بنك أسئلة الوحدة 11 (خيار من متعدد)', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    with open(input_txt, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by double newline to get individual question blocks
    question_blocks = content.strip().split('\n\n')

    for block in question_blocks:
        lines = block.strip().split('\n')
        if not lines:
            continue

        for line in lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # We can try to set RTL on paragraph runs, but setting alignment is usually enough for readability
            for run in p.runs:
                run.font.name = 'Arial'

        doc.add_paragraph() # Add empty line between questions

    doc.save(output_docx)
    print(f"Document saved to {output_docx}")

if __name__ == "__main__":
    create_word_document("questions_unit11_raw.txt", "questions_bank_unit11.docx")
