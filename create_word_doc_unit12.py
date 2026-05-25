from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re

def create_word_doc(input_file, output_file):
    doc = Document()

    # Set default style to RTL and Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Add title
    title = doc.add_paragraph('أسئلة اختيار من متعدد - الفصل 12')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)

    doc.add_paragraph() # Add some space

    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Split into individual questions
        questions = content.split('\n\n')

        for q in questions:
            if not q.strip():
                continue

            lines = q.strip().split('\n')
            for i, line in enumerate(lines):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # Optional: set text direction RTL if python-docx supports it easily
                # p.style.font.rtl = True

                run = p.add_run(line)
                if i == 0 or line.startswith('الإجابة الصحيحة'):
                    run.bold = True

            doc.add_paragraph() # Empty line between questions

        doc.save(output_file)
        print(f"Successfully created {output_file}")
    except Exception as e:
        print(f"Error creating Word document: {e}")

if __name__ == '__main__':
    create_word_doc('questions_unit12_raw.txt', 'questions_bank_unit12.docx')
